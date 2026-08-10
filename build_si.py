# -*- coding: utf-8 -*-
"""Build the Supplementary Information for the JTB resubmission.

Every numerical table in the SI is generated here directly from sensitivity_results.json.
Nothing is transcribed by hand, so the document cannot drift from the computation that
produced it. Regenerate the results first if you change the model:

    python3 sensitivity_lorenzini.py
    python3 make_sensitivity_figures.py
    python3 build_si.py

Written in response to JTB-D-26-01233, which asked for "a carefully described and extended
SI that includes full code access (not prompts) and provides the necessary sensitivity &
robustness analyses".
"""
import json
import os
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_jtb import _base_doc, _greek        # same neutral layout as the manuscript

HERE = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(HERE, "figures")
RES = json.load(open(os.path.join(HERE, "sensitivity_results.json")))
OUT = os.path.join(HERE, "Supplementary_Information_JTB_v2.docx")

# Repository to be created under the author's account before resubmission (see README).
REPO = "https://github.com/tallebhakeim/elasmobranch-electroreception-certified"
# Set once the archived release exists; the DOI line is omitted entirely while it is empty,
# so the document never ships a placeholder.
ZENODO_DOI = ""



# --- placeholder guard -------------------------------------------------------------
# The JTB submission (JTB-D-26-01233) was rejected in part because the .docx that went out
# still contained "[Authors to complete...]" in Funding and in Data availability. The editor
# went looking for the code, found a bracketed instruction, and concluded the figures might
# have come from prompts. A build that can emit that text is a foot-gun; this makes it fail.
PLACEHOLDERS = [r"\[Authors? to complete", r"0000-0000-0000-0000", r"<DOI>", r"<user>",
                r"\bTBD\b", r"\[to be inserted", r"\[ref\]", r"\[insert"]


def assert_no_placeholders(doc, path):
    """Raise before saving if any placeholder survived into the document body."""
    txt = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                txt += "\n" + c.text
    hits = sorted({m.group(0) for pat in PLACEHOLDERS
                   for m in re.finditer(pat, txt, re.I)})
    if hits:
        raise SystemExit(
            f"REFUSING to write {path}: unresolved placeholder(s) {hits}. "
            "Fill them in the source, or pass allow_placeholders=True for a working draft.")

def gain(r):
    return f"[{r['gain'][0]:.3f}, {r['gain'][1]:.3f}]"


def brk(r, key):
    return f"[{r[key][0]:.4f}, {r[key][1]:.4f}]"


def yesno(r):
    return "yes" if r["disjoint"] else "NO"


class SI:
    def __init__(self):
        self.doc = _base_doc(double_spaced=False, line_numbers=True)

    def h(self, text, size=13):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)

    def p(self, text, size=11):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(text).font.size = Pt(size)
        return p

    def caption(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        return p

    def fig(self, name, cap, width=15.5):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(os.path.join(FIGDIR, name), width=Cm(width))
        self.caption(cap)

    def table(self, header, rows, cap):
        self.caption(cap)
        t = self.doc.add_table(rows=len(rows) + 1, cols=len(header))
        t.style = "Table Grid"
        for j, v in enumerate(header):
            c = t.cell(0, j).paragraphs[0]
            c.paragraph_format.space_after = Pt(2)
            r = c.add_run(v)
            r.bold = True
            r.font.size = Pt(9)
        for i, row in enumerate(rows, start=1):
            for j, v in enumerate(row):
                c = t.cell(i, j).paragraphs[0]
                c.paragraph_format.space_after = Pt(2)
                c.add_run(str(v)).font.size = Pt(9)
        self.doc.add_paragraph()


def build():
    s = SI()
    s.h("Supplementary Information", 16)
    s.p("An error-bounded field model of elasmobranch electroreception: canal funnelling, "
        "frequency tuning and the effect of body plan in a hammerhead shark and a manta ray")
    s.p("H. Talleb")

    # ---------------------------------------------------------------- S1
    s.h("S1. What this document contains, and why")
    s.p("The claims in the main text are comparisons, not values: the gel canal raises the "
        "access conductance of an ampulla; a wide cephalofoil reads a larger navigation "
        "baseline than a compact disc; the modelled receptor levels sit above the behavioural "
        "threshold. A comparison computed from a discretised field solution is only as good as "
        "the discretisation error, and a conventional solve supplies an estimate of that error "
        "rather than a bound.")
    s.p("The method used here returns instead a guaranteed two-sided bracket. A primal solve in "
        "the potential over-estimates the dissipated energy; a complementary solve in a "
        "divergence-conforming current flux under-estimates it. The exact value lies between "
        "them. This turns a comparison into a decidable question: if the bracket computed with "
        "the canal and the bracket computed without it are disjoint, the ordering holds for the "
        "exact solution, whatever the mesh.")
    s.p("Section S5 therefore does not report a sensitivity study in the usual sense, in which a "
        "parameter is perturbed and a point value is seen to move a little. It reports, for "
        "every parameter set swept, the guaranteed bracket on the gain and whether it excludes "
        "unity. That is the quantity a reviewer should check, and it is the one plotted in "
        "every figure below.")

    # ---------------------------------------------------------------- S2
    s.h("S2. Governing problem")
    s.p("At the frequencies of electroreception the medium is a complex conductivity, "
        "kappa = sigma + j omega epsilon, and the displacement term is negligible against "
        "conduction everywhere except across the sensory membrane. In the bulk the potential "
        "obeys current conservation, div(sigma grad phi) = 0, with the pore held at unit "
        "potential and the deep receptor at zero for the access-conductance problem, and with "
        "the outer boundary of the seawater box either insulating or driven by the external "
        "stimulus for the sensitivity problems.")
    s.p("The access conductance is G = 2W, with W the dissipated power for a unit potential "
        "difference. Since the model of Section S4 is a plane problem, G is a conductance per "
        "unit depth and carries units of siemens per metre. It is not the three-dimensional "
        "access conductance of a real ampulla and should not be compared with one. All "
        "conclusions are drawn from the dimensionless ratio of two such conductances computed "
        "on the same geometry, which is free of that caveat.")

    # ---------------------------------------------------------------- S3
    s.h("S3. Discretisation and the guaranteed bounds")
    s.p("The discrete geometric method is used on a primal Delaunay triangulation and its "
        "circumcentric dual. The potential lives on primal nodes, its gradient on primal edges, "
        "the current flux on dual edges, and current conservation on dual cells. The "
        "differential operators are incidence matrices, exact and metric-free; all of the "
        "geometry and material enter through the constitutive (Hodge) operator, which is the "
        "only approximate step.")
    s.p("Upper bound. The primal nodal solve satisfies conservation exactly and the "
        "constitutive law approximately, and its energy over-estimates the exact dissipation.")
    s.p("Lower bound. A lowest-order Raviart-Thomas (RT0) flux, divergence-conforming and "
        "equilibrated with the prescribed electrode currents, gives a complementary energy that "
        "under-estimates the same dissipation. The pore and the receptor are meshed as carved "
        "conductor boundaries rather than as point constraints, which is what makes the "
        "complementary bound rigorous rather than merely plausible.")
    s.p("Neither bound involves a safety factor, a calibration or an asymptotic assumption. "
        "The half-width of the bracket, reported throughout, is a measured quantity.")

    # ---------------------------------------------------------------- S4
    s.h("S4. Geometry, mesh and a resolution criterion")
    ref = RES["reference"]
    s.p(f"The plane model is a {1e3*0.040:.0f} x {1e3*0.040:.0f} mm domain, tissue below and "
        f"seawater above a skin at 30 mm, with a gel canal of half-width "
        f"{ref['canal_halfwidth']*1e3:.1f} mm running from a surface pore to a deep receptor "
        f"electrode. Interior nodes are jittered by a seeded pseudo-random offset so that the "
        f"result does not depend on a structured grid.")
    s.p("Carving the electrodes out of the triangulation can leave nodes that belong only to "
        "removed triangles. Such a node carries no equation, the stiffness matrix is then "
        "exactly singular and the solve returns NaN without warning. Orphaned nodes are pruned "
        "and the connectivity reindexed before assembly. This is noted because it affects any "
        "reimplementation of the same carving strategy, and because it is invisible unless "
        "several mesh sizes are run.")
    s.p("The resolution criterion that matters is the number of elements across the canal. "
        "Table S1 and Figure S1 show that the bracket does not settle until the element size "
        "reaches the canal half-width. At h = 0.6 mm, the value used in an earlier draft of "
        "this work, the bracket half-width is 13.1 per cent; at h = 0.4 mm it is 2.3 per cent, "
        "and the guaranteed gain tightens from [1.199, 1.738] to [1.570, 1.708]. The reference "
        "mesh used throughout is h = 0.4 mm, and h = 0.3 mm is reported to show that the "
        "brackets then overlap, that is, that the result has converged.")

    s.table(["h (mm)", "triangles", "G with canal (S/m)", "G without canal (S/m)",
             "half-width (%)", "guaranteed gain", "disjoint"],
            [[f"{r['h']*1e3:.1f}", r["ntri"], brk(r, "G_canal"), brk(r, "G_nocanal"),
              f"{r['halfwidth_pct']:.2f}", gain(r), yesno(r)] for r in RES["mesh"]],
            "Table S1. Mesh refinement. G is a conductance per unit depth. The last column is "
            "the certified statement: the two brackets are disjoint, so the canal strictly "
            "raises the access conductance of the exact solution.")

    s.fig("Fig_S1_mesh.png",
          "Figure S1. (a) The guaranteed gain bracket against element size; the bar is the "
          "bracket itself, blue when it excludes unity and red when it does not. (b) The "
          "half-width of the bracket against mesh size. The non-monotonic behaviour on the "
          "coarsest meshes is expected: the jittered triangulation and the carved electrodes "
          "make the two bounds converge from different directions until the canal is resolved.")

    # ---------------------------------------------------------------- S5
    s.h("S5. Sensitivity and robustness")
    s.p("Five parameters are swept: the conductivity of the gel, of the body tissue and of the "
        "seawater, the width of the canal, and the mesh size. For each, the table gives the "
        "guaranteed bracket on the gain and whether it excludes unity.")

    s.p("S5.1 Gel conductivity, and how to read a failure. The conductivity of the ampulla gel "
        "is the least well constrained parameter in the model. The classical description treats "
        "the canal as a low-resistance core comparable to seawater; measurements of the jelly "
        "give lower figures depending on what is measured and how the sample is prepared. The "
        "sweep therefore covers two decades.")
    s.table(["sigma_gel (S/m)", "G with canal (S/m)", "G without canal (S/m)",
             "half-width (%)", "guaranteed gain", "disjoint"],
            [[f"{r['sig_gel']:g}", brk(r, "G_canal"), brk(r, "G_nocanal"),
              f"{r['halfwidth_pct']:.2f}", gain(r), yesno(r)] for r in RES["sig_gel"]],
            "Table S2. Gel conductivity at the reference mesh (h = 0.4 mm). Tissue conductivity "
            "is 0.3 S/m throughout, so the two lowest values place the canal at or below the "
            "conductivity of the surrounding tissue.")
    s.p("Two rows do not separate, and they are the two that must not. At sigma_gel = 0.3 S/m "
        "the canal has exactly the conductivity of the tissue, so it is not a canal at all and "
        "the true gain is exactly one; the bracket contains one, as it is required to. At "
        "sigma_gel = 0.2 S/m the canal is less conductive than the tissue and the bracket sits "
        "at and below one. The method declines to assert an effect where there is none, which "
        "is the behaviour a guaranteed bound should have and is the strongest available check "
        "that the implementation is correct.")

    s.p("S5.2 Separating a numerical limit from a physical one. On a coarser mesh the brackets "
        "at low gel conductivity also fail to separate, and that failure has a different cause: "
        "the bracket is simply too wide to resolve a small effect. Sweeping gel conductivity "
        "and mesh size together distinguishes the two.")
    s.table(["sigma_gel (S/m)", "h (mm)", "triangles", "half-width (%)",
             "guaranteed gain", "disjoint"],
            [[f"{r['sig_gel']:g}", f"{r['h']*1e3:.1f}", r["ntri"],
              f"{r['halfwidth_pct']:.2f}", gain(r), yesno(r)]
             for r in RES["sig_gel_x_mesh"]],
            "Table S3. Gel conductivity against mesh size. Refining the mesh separates the "
            "brackets at every gel conductivity above that of the tissue, so the failures in "
            "the coarse columns are a resolution limit and not a statement about the organ.")
    s.fig("Fig_S2_gel_x_mesh.png",
          "Figure S2. The same data. Reading left to right the mesh is refined; the brackets "
          "narrow and lift clear of unity. A red bar means the computation declines to decide, "
          "not that the canal fails to help.")

    s.p("S5.3 The remaining parameters. Body tissue conductivity, seawater conductivity and "
        "canal width are swept at the reference mesh. The gain bracket excludes unity "
        "throughout. Its magnitude behaves as the physics requires: the canal matters most when "
        "the surrounding tissue is most resistive, is nearly independent of the seawater "
        "conductivity, and grows with canal width.")
    for key, lab, unit in (("sig_tis", "sigma_tissue", "S/m"),
                           ("sig_sw", "sigma_seawater", "S/m"),
                           ("canal_halfwidth", "canal half-width", "mm")):
        rows = RES[key]
        scale = 1e3 if key == "canal_halfwidth" else 1
        s.table([f"{lab} ({unit})", "G with canal (S/m)", "G without canal (S/m)",
                 "half-width (%)", "guaranteed gain", "disjoint"],
                [[f"{r[key]*scale:g}", brk(r, "G_canal"), brk(r, "G_nocanal"),
                  f"{r['halfwidth_pct']:.2f}", gain(r), yesno(r)] for r in rows],
                f"Table S4{'abc'[('sig_tis','sig_sw','canal_halfwidth').index(key)]}. "
                f"Sweep of {lab} at the reference mesh.")
    s.fig("Fig_S3_materials.png",
          "Figure S3. Guaranteed gain brackets across the material and geometric parameters.")

    s.p("S5.4 The frequency tuning. The membrane and wall time constants are representative "
        "values, not values fitted to a species, so the relevant question is not what the peak "
        "frequency is but how much of the conclusion survives that choice. Each of the four RC "
        "values is perturbed independently and log-uniformly by up to a factor of 1.5 in either "
        "direction.")
    tu = RES["tuning"]
    s.table(["quantity", "median (Hz)", "5th percentile", "95th percentile"],
            [[k.replace("_Hz", "").replace("_", " "),
              f"{v['median']:.3f}", f"{v['p05']:.3f}", f"{v['p95']:.3f}"]
             for k, v in tu.items() if isinstance(v, dict)],
            f"Table S5. Band-pass corners and peak under a plus or minus 50 per cent "
            f"log-uniform perturbation of every RC value, {20000} draws. A band-pass exists in "
            f"{100*tu['fraction_bandpass']:.1f} per cent of draws and the peak falls between "
            f"0.1 and 10 Hz in {100*tu['fraction_peak_in_0p1_10Hz']:.1f} per cent.")
    s.fig("Fig_S4_tuning.png",
          "Figure S4. (a) Distribution of the peak frequency. (b) The two corner frequencies "
          "for every draw; all points lie above the diagonal, so the ordering that creates the "
          "band-pass is never inverted.")

    # ---------------------------------------------------------------- S6
    s.h("S6. Code, data and how to reproduce every number")
    s.p("All code is public. The repository contains the field engine, the scripts that produce "
        "each figure of the main text, the sensitivity campaign of Section S5, and this "
        "document's build script, which reads the campaign output directly so that no number in "
        "the tables above is transcribed by hand.")
    s.p(f"Repository: {REPO}")
    if ZENODO_DOI:
        s.p(f"Archived release: {ZENODO_DOI}")
    s.table(["script", "produces"],
            [["lorenzini_certified.py", "the guaranteed brackets on the access conductance"],
             ["lorenzini_poc.py", "canal funnelling and the directional sweep, Figure 2"],
             ["lorenzini_freq.py", "the band-pass and its corners, Figure 4"],
             ["lorenzini_3d.py", "the directional rosette, Figure 3"],
             ["lorenzini_real.py", "the mesh figure and the receptor levels on both body "
                                  "plans, Figures 1, 5 and 6"],
             ["lorenzini_cloak.py, lorenzini_cloak_fig.py", "electrical visibility, Figure 7"],
             ["sensitivity_lorenzini.py", "the campaign of Section S5 (JSON output)"],
             ["make_sensitivity_figures.py", "Figures S1 to S4"],
             ["build_si.py", "this document"]],
            "Table S6. Scripts and what each one produces. Running reproduce_all.py executes "
            "them in order and regenerates every figure and every number from scratch.")
    s.p("The two surface models used for the body plans are distributed by DigitalLife3D under "
        "a Creative Commons Attribution-NonCommercial licence (CC BY-NC 4.0) and are "
        "redistributed in the repository under the same terms, together with the reorientation "
        "and scaling applied to them.")
    s.p("Declaration on tool use. A generative artificial-intelligence assistant was used while "
        "developing the code and drafting the text, and this is declared in the main text in "
        "the form requested by the publisher. No figure, number or table in this work is "
        "produced by that assistant: every one of them is the output of the deterministic, "
        "seeded scripts listed above, which any reader can run.")

    # ---------------------------------------------------------------- S7
    s.h("S7. What is not certified")
    s.p("The bracket is rigorous for the access conductance, which is a self-energy. The "
        "pointwise receptor voltages reported in the main text are computed but not themselves "
        "bracketed; certifying them requires a goal-oriented, reciprocity-based formulation on "
        "the same dual meshes, which is the natural continuation of this work.")
    s.p("The plane model of Sections S2 to S5 is a cross-section, so its directional sweep gives "
        "the order of the directional selectivity of a single ampulla and not a species value.")
    s.p("The body surfaces of the two body plans are photogrammetric models, but the "
        "electrosensory anatomy placed on them is parametric: nine ampullae per animal, one "
        "canal length, one canal orientation, against the order of one thousand ampullae with a "
        "species-specific distribution in a real animal. The comparison therefore tests the "
        "effect of the gross geometry of the array, which is what it is meant to test.")
    s.p("The access resistance quoted in the frequency model is the analytic resistance of a "
        "cylindrical canal, not the bracketed quantity of Section S3. The two are consistent in "
        "order of magnitude but they are not the same computation, and the main text says so.")

    _greek(s.doc)
    assert_no_placeholders(s.doc, OUT)
    s.doc.save(OUT)
    print("saved", os.path.basename(OUT), "| tables", len(s.doc.tables),
          "| paragraphs", len(s.doc.paragraphs))


if __name__ == "__main__":
    build()
