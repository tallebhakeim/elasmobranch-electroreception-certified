# -*- coding: utf-8 -*-
"""Build the Journal of Theoretical Biology submission package for the Lorenzini
electroreception article (Article 7, retargeted from J. R. Soc. Interface).

Produces three files, which is what the Elsevier editorial system expects:
  Article_Lorenzini_Electroreception_JTB_v1.docx   the manuscript
  Highlights_Lorenzini_JTB_v1.docx                 the required Highlights file
  Cover_Letter_Lorenzini_JTB_v1.docx               the cover letter

Neutral single-column layout (Elsevier "your paper, your way"), double spacing and
continuous line numbers for the reviewers. Equations are real OMML, converted from
LaTeX through pandoc, so they are editable in Word.

Run from this directory:  python3 build_jtb.py
Figures are read from FIGDIR (see below).
"""
import os
import re
import zipfile

import pypandoc
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "ARTICLE_LORENZINI_JTB_EN.md")
FIGDIR = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "Article_Lorenzini_Electroreception_JTB_v1.docx")
OUT_HL = os.path.join(HERE, "Highlights_Lorenzini_JTB_v1.docx")
OUT_CL = os.path.join(HERE, "Cover_Letter_Lorenzini_JTB_v1.docx")

# ---------------------------------------------------------------- equations
EQ_LATEX = {
    1: r"\kappa = \sigma + j\omega\varepsilon",
    2: r"W_{\mathrm{lower}} \le W_{\mathrm{exact}} \le W_{\mathrm{upper}}",
    3: r"Y_m = \frac{\sigma_m + j\omega\varepsilon_m}{d_m}",
    4: r"H(\omega) = \frac{j\omega\tau_w}{1+j\omega\tau_w}\cdot\frac{1}{1+j\omega\tau_m}",
    5: r"f_0 = \frac{1}{2\pi\sqrt{\tau_w\tau_m}}",
    6: r"\mathbf{E}_{\mathrm{mot}} = \mathbf{v}\times\mathbf{B}",
    7: r"\alpha = \frac{\sigma_b-\sigma_{sw}}{\sigma_b+2\sigma_{sw}}",
    8: r"\sigma_{\mathrm{eff}}(\sigma_{\mathrm{core}},\sigma_{\mathrm{shell}},f) = \sigma_{sw}",
}

# ------------------------------------------------- figures (citation order)
FIG = {
    1: "Fig_mesh3d.png",
    2: "Fig_lorenzini.png",
    3: "Fig_lorenzini_3d.png",
    4: "Fig_lorenzini_freq.png",
    5: "Fig_shark_EB.png",
    6: "Fig_ray_vs_shark.png",
    7: "Fig_cloak.png",
}
CAP = {
    1: "Figure 1. The tetrahedral mesh behind the solve, on a great hammerhead head "
       "(Sphyrna mokarran). (a) The body surface immersed in seawater with the ventral "
       "cephalofoil array, pores and gel canals. (b) A horizontal slab of the mesh, "
       "coloured by material, resolving tissue, seawater and the gel canals. Surface "
       "model: DigitalLife3D (CC BY-NC 4.0).",
    2: "Figure 2. Funnelling and directionality of one ampulla. (a) Under an external "
       "field the gel canal carries the surface potential to the receptor while the "
       "resistive body stays nearly equipotential. (b) The receptor voltage varies by a "
       "factor of six with the direction of the applied field.",
    3: "Figure 3. Array on a three-dimensional head: each ampulla canal points along a "
       "different azimuth and each is most sensitive to the field direction along its own "
       "canal, so the population response encodes field direction (the directional rosette).",
    4: "Figure 4. Frequency tuning from the capacitive sensory membrane. (a) The response "
       "is a band-pass peaking near 1.3 Hz, from the membrane low-pass and the skin and "
       "canal-wall high-pass, with the transfer function and corner frequencies of "
       "equations (4) and (5). (b) The membrane capacitance sets the upper corner, so the "
       "passband is tunable by that one parameter.",
    5: "Figure 5. Electric and magnetic levels on the hammerhead. (a) Receptor levels per "
       "ampulla across the cephalofoil; both modalities sit above the behavioural "
       "threshold. (b) The wide cephalofoil spreads the ventral array, so the uniform "
       "navigation field gives a V-shaped lateral baseline, smallest at the centre and "
       "largest at the wing tips, while a prey dipole gives a localised peak.",
    6: "Figure 6. Hammerhead versus manta on photogrammetric surface models (DigitalLife3D, "
       "CC BY-NC 4.0). (a, b) Receptor levels per ampulla for each. (c, d) Ventral views "
       "with ampullae sized by the prey response: the wide cephalofoil spreads the array "
       "over a large lateral span and reads the largest navigation baseline, while the "
       "manta array clusters ventrally around the mouth.",
    7: "Figure 7. Electrical visibility of a body in seawater. (a) The perturbation "
       "potential of a bare body is a dipole the sense reads. (b) A neutral-inclusion "
       "coating collapses it. (c) Sweeping the coating conductivity gives a deep null; an "
       "insulating coating (far left) is worse than bare skin. (d) Only a conductive, "
       "water-matched coating suppresses the signature.",
}

# ------------------------------------------------------------------ table 1
TABLE1 = [
    ("Quantity", "Symbol", "Value", "Source"),
    ("Seawater conductivity", "sigma_sw", "4 S/m", "standard"),
    ("Body tissue conductivity", "sigma_tis", "0.3 S/m", "standard"),
    ("Ampulla gel conductivity", "sigma_gel", "4 S/m", "assumed as seawater"),
    ("Canal length", "L", "50 mm", "representative"),
    ("Canal radius", "r", "0.5 mm", "representative"),
    ("Canal access resistance", "R_a", "15.9 kOhm", "analytic cylinder, section 2.3"),
    ("Sensory membrane resistance", "R_m", "6.6 MOhm", "representative"),
    ("Sensory membrane capacitance", "C_m", "3.0 nF", "representative"),
    ("Skin and canal-wall resistance", "R_w", "2.0 MOhm", "representative"),
    ("Skin and canal-wall capacitance", "C_w", "0.4 uF", "representative"),
    ("High-pass corner", "f_1", "0.20 Hz", "derived"),
    ("Low-pass corner", "f_2", "8.0 Hz", "derived"),
    ("Peak of the passband", "f_0", "1.3 Hz", "derived, eq. (5)"),
    ("Swimming speed", "v", "1 m/s", "representative"),
    ("Geomagnetic flux density", "B", "50 uT", "standard"),
    ("Motional field", "E_mot", "50 uV/m", "derived, eq. (6)"),
    ("Behavioural threshold", "-", "a few nV/cm", "literature"),
    ("Body length of the scaled models", "L_body", "1 m", "model"),
    ("Ampullae per animal", "n", "9", "model"),
    ("Reference element size", "h", "0.4 mm", "resolves the canal, see SI S4"),
]

HIGHLIGHTS = [
    "One field model gives the electric and the magnetic sensitivity of an ampulla array.",
    "Complementary energy bounds bracket the ampulla access conductance to 2.3 per cent.",
    "The gel canal raises it by a guaranteed factor between 1.57 and 1.71.",
    "A capacitive sensory membrane reproduces the measured band-pass near 1.3 hertz.",
    "Hammerhead and manta body plans read the same fields with different geometries.",
]

# Filled in by the authors before submission; the ORCID line is omitted while empty, so the
# document never ships a 0000-0000-0000-0000 placeholder.
ORCIDS = {"H. Talleb": "0000-0002-2734-6242"}

GREEK = [(r'(?<![A-Za-z])epsilon(?![A-Za-z])', 'ε'), (r'(?<![A-Za-z])sigma(?![A-Za-z])', 'σ'),
         (r'(?<![A-Za-z])omega(?![A-Za-z])', 'ω'), (r'(?<![A-Za-z])kappa(?![A-Za-z])', 'κ'),
         (r'(?<![A-Za-z])tau(?![A-Za-z])', 'τ'), (r'(?<![A-Za-z])alpha(?![A-Za-z])', 'α'),
         (r'(?<![A-Za-z])rho(?![A-Za-z])', 'ρ'), (r'(?<![A-Za-z])mu(?![A-Za-z])', 'μ'),
         ('<=', '≤'), ('Ohm', 'Ω'), (r'(?<![A-Za-z])uF(?![A-Za-z])', 'μF'),
         (r'(?<![A-Za-z])uT(?![A-Za-z])', 'μT'), (r'(?<![A-Za-z])uV(?![A-Za-z])', 'μV')]



# --- placeholder guard -------------------------------------------------------------
# The JTB submission (JTB-D-26-01233) was rejected in part because the .docx that went out
# still contained "[Authors to complete...]" in Funding and in Data availability. The editor
# went looking for the code, found a bracketed instruction, and concluded the figures might
# have come from prompts. A build that can emit that text is a foot-gun; this makes it fail.
PLACEHOLDERS = [r"\[Authors? to complete", r"0000-0000-0000-0000", r"<DOI>", r"<user>",
                r"\bTBD\b", r"\[to be inserted", r"\[ref\]", r"\[insert"]


def assert_no_placeholders(doc, path):
    """Raise before saving if any placeholder survived into the document body."""
    txt = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                txt += "\n" + c.text
    hits = sorted({m.group(0) for pat in PLACEHOLDERS
                   for m in re.finditer(pat, txt, re.I)})
    if hits:
        raise SystemExit(
            f"REFUSING to write {path}: unresolved placeholder(s) {hits}. "
            "Fill them in the source, or pass allow_placeholders=True for a working draft.")

def _omml(eq_latex):
    """LaTeX -> OMML, via a throwaway pandoc docx."""
    tmp = os.path.join(HERE, "_eqs.docx")
    md = "\n\n".join(f"$${eq_latex[k]}$$" for k in sorted(eq_latex))
    pypandoc.convert_text(md, "docx", format="markdown", outputfile=tmp)
    x = zipfile.ZipFile(tmp).read("word/document.xml").decode("utf8")
    raw = re.findall(r"<m:oMath>.*?</m:oMath>", x, re.S)
    return {k: f.replace("<m:oMath>", f'<m:oMath xmlns:m="{_M}" xmlns:w="{_W}">', 1)
            for k, f in zip(sorted(eq_latex), raw)}


def _base_doc(double_spaced=True, line_numbers=True):
    """A neutral single-column document: Times New Roman 12, wide margins."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    if double_spaced:
        pf.line_spacing = 2.0
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.5)
    if line_numbers:
        sec._sectPr.append(parse_xml(
            f'<w:lnNumType xmlns:w="{_W}" w:countBy="1" w:restart="continuous" w:distance="360"/>'))
    return doc


def _greek(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            t = r.text
            if not t:
                continue
            for pat, rep in GREEK:
                t = re.sub(pat, rep, t)
            if t != r.text:
                r.text = t
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        t = r.text
                        for pat, rep in GREEK:
                            t = re.sub(pat, rep, t)
                        if t != r.text:
                            r.text = t


def _check_figures():
    """Fail early and legibly if a figure is absent, rather than deep inside python-docx.
    Three of the seven figures need the DigitalLife3D surface models; a fresh clone without
    them cannot build the manuscript, and should say so plainly."""
    missing = [(n, f) for n, f in sorted(FIG.items())
               if not os.path.exists(os.path.join(FIGDIR, f))]
    if missing:
        raise SystemExit(
            "cannot build the manuscript: missing figure(s) "
            + ", ".join(f"{f} (Figure {n})" for n, f in missing)
            + f"\n  looked in {FIGDIR}"
            + "\n  Figures 1, 5 and 6 need the DigitalLife3D surface models; see"
              " meshes/README.md, then rerun reproduce_all.py.")


def build_manuscript():
    _check_figures()
    EQ = _omml(EQ_LATEX)
    doc = _base_doc()

    def para(text, size=12, bold=False, italic=False, align=None, space_before=0, spacing=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        if spacing is not None:
            p.paragraph_format.line_spacing = spacing
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return p

    def addfig(n):
        path = os.path.join(FIGDIR, FIG[n])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.add_run().add_picture(path, width=Cm(15.0))
        c = para(CAP[n], size=10, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0)
        c.paragraph_format.space_after = Pt(14)

    def addtable():
        c = para("Table 1. Parameters of the model. Values marked representative are "
                 "chosen within the range reported in the literature and are not fitted "
                 "to a species.", size=10, spacing=1.0)
        c.paragraph_format.space_after = Pt(6)
        t = doc.add_table(rows=len(TABLE1), cols=4)
        t.style = "Table Grid"
        for i, row in enumerate(TABLE1):
            for j, val in enumerate(row):
                cell = t.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(val)
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"
                r.bold = (i == 0)
        doc.add_paragraph()

    def addeq(line):
        n = int(re.search(r'\((\d+)\)\s*$', line).group(1))
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        if n in EQ:
            p._p.append(parse_xml(EQ[n]))
        else:
            p.add_run(re.sub(r'\s*\(\d+\)\s*$', '', line)).italic = True
        p.add_run("\t(" + str(n) + ")")

    text = open(SRC, encoding="utf8").read()
    title = next(l[2:].strip() for l in text.split("\n") if l.startswith("# "))
    pre, _, rest = text.partition("## Abstract")
    ab, _, after = rest.partition("## 1. Introduction")
    meta = [l.strip() for l in pre.split("\n") if l.strip() and not l.startswith("# ")]

    para(title, size=16, bold=True, spacing=1.2)
    for m in meta:
        if m.startswith("Keywords:"):
            continue
        para(m, size=11, spacing=1.2)
    para("Abstract", size=12, bold=True, space_before=12)
    for p_ in [q.strip() for q in ab.strip().split("\n") if q.strip()]:
        para(p_)
    if ORCIDS:
        para("ORCID: " + "; ".join(f"{k}, {v}" for k, v in ORCIDS.items()), size=11, spacing=1.2)
    kw = [m for m in meta if m.startswith("Keywords:")]
    if kw:
        para(kw[0], size=11, spacing=1.2)

    pend = None
    inref = False
    for raw in ("## 1. Introduction" + after).split("\n"):
        l = raw.rstrip()
        if not l.strip():
            continue
        if l.startswith("## References"):
            if pend:
                addfig(pend)
                pend = None
            inref = True
            para("References", size=13, bold=True, space_before=12)
            continue
        if inref:
            p = para(l.strip(), size=11, spacing=1.0)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            continue
        if l.startswith("### ") or l.startswith("## "):
            if pend:
                addfig(pend)
                pend = None
            lvl = 3 if l.startswith("### ") else 2
            h = l[lvl + 1:].strip()
            mf = re.search(r'\s*\(figure (\d+)\)', h)
            if mf:
                pend = int(mf.group(1))
                h = h[:mf.start()].rstrip()
            para(h, size=14 if lvl == 2 else 12, bold=True, space_before=12, spacing=1.2)
            continue
        if l.strip().startswith("Table 1."):
            addtable()
            continue
        if re.search(r'\(\d+\)\s*$', l.strip()) and "=" in l:
            addeq(l.strip())
            continue
        para(l.strip())
    if pend:
        addfig(pend)

    _greek(doc)
    assert_no_placeholders(doc, OUT)
    doc.save(OUT)
    print("saved", os.path.basename(OUT), "| paragraphs", len(doc.paragraphs))


def build_highlights():
    doc = _base_doc(double_spaced=False, line_numbers=False)
    p = doc.add_paragraph()
    r = p.add_run("Highlights")
    r.bold = True
    r.font.size = Pt(14)
    p = doc.add_paragraph()
    r = p.add_run("An error-bounded field model of elasmobranch electroreception: canal "
                  "funnelling, frequency tuning and the effect of body plan in a hammerhead "
                  "shark and a manta ray")
    r.italic = True
    r.font.size = Pt(11)
    doc.add_paragraph()
    for h in HIGHLIGHTS:
        assert len(h) <= 85, f"highlight too long ({len(h)} chars): {h}"
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(h).font.size = Pt(12)
    assert_no_placeholders(doc, OUT_HL)
    doc.save(OUT_HL)
    print("saved", os.path.basename(OUT_HL), "|",
          ", ".join(str(len(h)) for h in HIGHLIGHTS), "chars")


COVER = """To the Editors
Journal of Theoretical Biology

Dear Editors,

I submit the manuscript "An error-bounded field model of elasmobranch electroreception: canal \
funnelling, frequency tuning and the effect of body plan in a hammerhead shark and a manta ray".

This is a new submission of work previously declined as JTB-D-26-01233. That decision asked for \
three things before the work could go to review: the code, a carefully described supplementary \
information, and sensitivity and robustness analyses. I have done all three, and the exercise \
changed the paper enough that it is worth saying how.

The code is public and the article now has a single entry point, reproduce_all.py, which \
regenerates every figure and every number from scratch in about two minutes and reports what \
ran, what was skipped and why. The supplementary information sets out the discretisation, the \
construction of the complementary bound, the mesh resolution criterion and the full campaign; \
every numerical table in it is generated directly from the campaign output rather than \
transcribed, so the document cannot drift from the computation behind it. I regret that the \
previous version left the reader no way to check any of this, and that its data availability \
section was unfinished when it went out.

On the sensitivity analyses, the method makes a stronger statement possible than the usual one. \
Because the solver returns a guaranteed two-sided bracket rather than a point value, a \
comparison can be settled outright: if the bracket computed with the gel canal and the bracket \
computed without it are disjoint, the canal provably raises the access conductance of the exact \
solution. I therefore report, for all thirty-five parameter sets swept, the guaranteed bracket \
on the gain and whether it excludes unity. Thirty-two do. The three that do not are the three \
that must not: when the gel is given the conductivity of the surrounding tissue the canal is no \
longer a canal, the true gain is exactly one, and the bracket contains one. The method declines \
to assert an effect where there is none.

Carrying out that work also exposed three errors in the previous version, which are corrected \
here. The access conductance was quoted in the wrong units, by a factor of one thousand and in \
the wrong dimension: the model is a plane problem, so the quantity is a conductance per unit \
depth. The mesh did not resolve the canal, which inflated the bracket to thirteen per cent; at a \
resolved mesh it is 2.3 per cent and the guaranteed gain tightens from a vague "about 1.5" to \
[1.57, 1.71]. And the access resistance used in the frequency model was attributed to the field \
solve when it is in fact the analytic resistance of a cylinder; the text now says so. I would \
not have found any of these without the analyses the decision letter asked for.

The biology is unchanged and is the reason for submitting here. One ampulla is already a \
directional sensor, so an array reads field direction. The capacitive sensory epithelium yields \
a band-pass peaking near 1.3 Hz, within the measured range, without that tuning being imposed, \
and it survives a wide perturbation of the parameters that produce it. On the two body plans, a \
uniform navigation field and a local prey dipole leave different spatial signatures on the same \
array, the first a lateral baseline across the cephalofoil and the second a localised peak, so \
the two are separable by the shape of the population response rather than by its amplitude. \
That is a prediction array-level electrophysiology can test, and it gives a quantitative form to \
the electrosensory argument for the hammerhead cephalofoil.

The manuscript is original, is not under consideration elsewhere, and I am its sole author. I \
declare no competing interests. The use of a generative artificial-intelligence assistant in \
developing the code and drafting the text is declared in the manuscript in the form requested by \
the publisher; no figure, number or table is produced by it, and the repository makes that \
checkable.

Thank you for reconsidering this work.

Yours sincerely,

Hakeim Talleb
Sorbonne Universite, GeePs (CNRS, CentraleSupelec, Universite Paris-Saclay)
hakeim.talleb@sorbonne-universite.fr
"""


def build_cover():
    doc = _base_doc(double_spaced=False, line_numbers=False)
    for block in COVER.strip().split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run(block.strip())
    assert_no_placeholders(doc, OUT_CL)
    doc.save(OUT_CL)
    print("saved", os.path.basename(OUT_CL))


if __name__ == "__main__":
    build_manuscript()
    build_highlights()
    build_cover()
